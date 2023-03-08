


#Upload File
from ..forms import UploadFileForm
from django.core.files.storage import default_storage
from django.shortcuts import render
from django.http import HttpResponse
from django.shortcuts import redirect
from ..models.models import Mcq,Subject,SubjectAccess
from ..ultils import isManger
from docx import Document
import os
import shutil
import time
import re
import easyocr
from PIL import Image

import torch
import numpy as np
from sentence_transformers import SentenceTransformer
from sentence_transformers import SentenceTransformer, util
from ..Dbcontext import connector
from.home_view import home
#regex find image_path 
img_regex = "(?<=\[file: ).+?(?=\])"
ques_regex = "(\[file:).*?(\])"
op_regex = "([a-zA-Z]\.)"

#Easy OCR 
reader = easyocr.Reader(['en','vi'])

print("Connecting to pinecone")
conn = connector()

device = 'cuda' if torch.cuda.is_available() else 'cpu'
print("Device:",device)
print("Loading SBert Model ")
SbertModel = SentenceTransformer('model\\all-mpnet-finetune-5epochs',device=device)

#Sbert 

def import_view(request):

    if request.method == "POST":

        if 'submit_import' in request.POST:
            checked_items = request.POST.getlist("check_import")
            #show import mcq
            print(checked_items)
            print("--------------------------")
            temp_dct = request.session['temp_mcq']

            currentUser = request.user

            for key_mcq in checked_items:
                mcq = temp_dct[key_mcq]
                encode = mcq["encode"]
                subject = mcq["subject"]
                #generate new id
                user_name = request.user.username
                new_id = subject+"-"+user_name+"-"+get_time()
                #get attr
                question = mcq["question"]
                options = mcq["options"]
                answer = mcq["answer"]
                haveImage = mcq["haveImage"]

                if haveImage:
                    images = new_id+".png"
                    temp_path = "media/temp/"+mcq["image"]
                    image_path = "media/images/"+images
                    new_path = "images/"+images
                    shutil.move(temp_path, image_path)

                    save_mcq= Mcq( 
                        qid =  new_id,
                        question = question,
                        options = options,
                        q_image = images,
                        answer_q = answer,
                        subject = subject,
                        contain_img = haveImage,
                        img_file = new_path ,
                        user = currentUser
                        )
                    # save_mcq.save()
                else:
                    save_mcq= Mcq( 
                        qid =  new_id,
                        question = question,
                        options = options,
                        answer_q = answer,
                        subject = subject,
                        contain_img = haveImage,
                        user = currentUser
                         )
                    images = ""
                    
                # save into database
                save_mcq.save()
                print("Import Mcq success !")
                print(save_mcq.subject)
                print(save_mcq.user)

                # save into Pinecone
                result = conn.upload(qid=new_id,encode=encode,question=question,contain=haveImage,q_image=images,subject=subject)
                print(result)
            # Remove file in temp 

            shutil.rmtree('media/temp')

            return redirect(home)
        else:
            
            #Import
            form = UploadFileForm(request.POST,request.FILES)
            files = request.FILES.getlist('file')
            list_docx={}
            list_images={} 
            importSubject = request.session['subjectImport']
            #Load Images and docx file  
            for file in files:
                name_file,ex = os.path.splitext(str(file))
                # document = Document(file)
                if ex in ['.docx']:
                    print("found document")
                    list_docx[str(file)] = []
                    doc = Document(file)
                    for table in doc.tables:
                        
                        new_mcq = row2mcq(table,importSubject)
                        list_docx[str(file)].append(new_mcq)
                    # print("Total questions:",len(doc.tables))
                else:
                    list_images[str(file)] = file
                    save_path = "temp/"+file.name
                    default_storage.save(save_path, file)
                    
                    #reading Url 
                    # file_url = default_storage.url(save_path)

                    # print("Image_path:",file_url)
            duplicate_mcq =[]
            num_dup = 0
            mcqs_set={}
            mcq_lst = []
            mcq_lst_encode = []
            for key in list_docx.keys():
                for mcq in list_docx[key]:
                    #store in session 
                   
                    print(mcq.qid)
                    mcq_image_text =""
                    
                    if mcq.contain_img:
                        file_img = list_images[mcq.q_image]
                        mcq_image_text = ocr2Text(file_img)
                    mcq_form = mcq.getMcq(mcq_image_text)

                        # print(mcq_form)
                    encode = SbertModel.encode(str(mcq_form)).tolist()
                    print(str(mcq.subject).strip())
                    result = conn.query_mcqs_encode(encode,str(mcq.subject).strip(),k=5)

                        # Filter dictionary by keeping elements whose keys are divisible by 2
                    list_id_dup = [ (d["id"],d["score"]) for d in result["matches"] if d["score"] >= 0.75] 
                    if len(list_id_dup) != 0:

                        duplicate_mcq.append((mcq,list_id_dup))
                        num_dup += len(list_id_dup)

                    mcqs_set[mcq.qid] = {"id":mcq.qid,"question":mcq.question,"image":mcq.q_image,"options":mcq.options,"answer":mcq.answer_q,"subject":mcq.subject,"haveImage":mcq.contain_img,"encode":encode,"qid":mcq.qid} 
                    mcq_lst.append(mcq.qid)
                    mcq_lst_encode.append(encode)
            

            pairs = cosin_pair(mcq_lst_encode)
            for pair in pairs:
                i, j = pair['index']
                temp_mcq = mcqs_set[mcq_lst[i]]
                # print("temp_mcq",temp_mcq)
                duplicate_mcq.append((temp_mcq,[(mcq_lst[j],pair['score'])]))
                num_dup += 1

                print("{} \t\t {} \t\t Score: {:.4f}".format(mcq_lst[i], mcq_lst[j], pair['score']))

            request.session['temp_mcq'] = mcqs_set
            print("import length:",len(mcqs_set))
            # print("import length:",mcqs_set)
     
            return render(request, 'import.html',{'dup_result':duplicate_mcq,"num_dup":num_dup})
    else:
         #authen the subject import 
        subjectImport = request.GET.get("subject")
        currentUser = request.user
        checkAuthen = authenUserSubject(subjectImport,currentUser)
        if checkAuthen:
            print("Ok !")
        else:
            return redirect('home')

        request.session['subjectImport'] =  subjectImport
        form = UploadFileForm()
    return render(request, 'import.html',{'form':form,'subjectImport':subjectImport})



def authenUserSubject(subject,Teacher):
    print("Import subject",subject)
    if subject == None:
        return False
    
    if isManger(Teacher):
        print("is manager !")
        return True
    try:
        ImportSubject = Subject.objects.get(subject=subject)
    except:
        print("Cant not find subject")
        return False
    
    try:
        print(Teacher.username)
        getAccess = SubjectAccess.objects.get(teacher = Teacher.username, subject = ImportSubject)
        return True
    except:
        print("Cant not find Access")
        return False 


def cosin_pair(encode):
  cosine_scores = util.cos_sim(encode, encode)
  pairs = []

  for i in range(len(cosine_scores)-1):
    for j in range(i+1, len(cosine_scores)):
        # Compare thresh hold 
        score = np.array(cosine_scores[i][j])
        if cosine_scores[i][j] >= 0.75:
            pairs.append({'index': [i, j], 'score': score})

  #Sort scores in decreasing order
  pairs = sorted(pairs, key=lambda x: x['score'], reverse=True)
  return pairs


def ocr2Text(file):
    path="temp.png"
    img = Image.open(file)
    img.save(path)
    img_text = reader.readtext(path, paragraph=True)

    question = ""
    for i in img_text:
        question += i[1] +" "
    return question

def get_time():
    t = time.time()
    t_s = int(t)
    return str(t_s)

def row2mcq(table,subject="math"):
    data = [[cell.text for cell in row.cells] for row in table.rows]
    mcq_options = []
    ans =""
    for row in data:
        if(row[0].strip() == "answer"):
            ans = row[1]
        if("QN" in row[0]):
            mcq_name = row[0]
            text = row[1]
            text = text.replace("\n"," ")
            path = re.findall(img_regex,text) 
            if(len(path) == 0 ):
                path =""
            else:
                path = path[0]                
            question = re.sub(ques_regex, '',text )
            mcq_question = question
            # print("Question:",question)
            mcq_images = path
       
        if(re.match(op_regex,row[0])):
            text = row[1]
            text = text.replace("\n"," ")
            if len(text.strip()) != 0 :
                mcq_options.append(text)
    # print(mcq_images)
    # print(mcq_options)
    contain_img = False if len(mcq_images.strip()) <=0 else True
    if contain_img:
        new_mcq = Mcq(qid = mcq_name,question = mcq_question,options = str(mcq_options) , q_image = mcq_images , answer_q = ans, subject = subject, contain_img = contain_img, img_file = "temp/"+  mcq_images )
    else:
        new_mcq = Mcq(qid = mcq_name,question = mcq_question,options = str(mcq_options) , q_image = mcq_images , answer_q = ans, subject = subject, contain_img = contain_img )
    return new_mcq




