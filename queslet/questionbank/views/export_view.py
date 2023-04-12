from django.shortcuts import render,redirect
from django.http import HttpRequest
from ..models.models import Mcq
from django.core.paginator import Paginator
from django.contrib import messages
from docx.shared import Inches, Cm
import docx
import os
from .home_view import get_context

from ..ultils import get_time


# Create your views here.

def export(request):

    if not request.user.is_authenticated:
        return redirect('login')
    

    # result = Mcq.objects.all()
    # have_img =Mcq.objects.filter(contain_img =True)
    # subjects = Mcq.objects.values('subject').distinct()


    # context ={"lst_mcqs":result,
    #               "total":len(result),
    #               "num_img":len(have_img),
    #               "num_subject":len(subjects),
    #     }   
    context = get_context(request=request)

    if request.method == "GET":
        # Load data
        #database 


        subjects_selected = request.GET.get("subject")
        print(subjects_selected)

        request.session['subjects_selected'] = subjects_selected
        
        context["subjects_selected"] = subjects_selected
        return render(request, 'export.html',context)
    elif request.method == 'POST':
        time = get_time()
        subjects_selected = request.POST.get("subjects_selected")
        id_list = request.POST.getlist('checkbox')
        print(id_list)
        username = request.user.username
        mydoc = docx.Document()
        for id in id_list:
            export_mcq = Mcq.objects.get(qid= id )
            # qid = export_mcq.qid
            # question = str(export_mcq.question)
            # options = str(export_mcq.options)
            # images = str(export_mcq.q_image)
            # new_mcq = mcq(qid,question=question,options=options,answer="",images=images)
            save_ques(mydoc,export_mcq,username ,subjects_selected ,time)



        messages.success(request, ("Export successfully ! "))
        return render(request, 'export.html',context)
    

def save_ques(mydoc,new_mcq,username,subject,time):
# model write to Word file 
#caculate row and columns 
    a = 97
    # chr(a + i)
    col = 2 
    row = 1
    options = str(new_mcq.options)

    question = "" if str(new_mcq.question) is None else str(new_mcq.question)

    list_op =[ o.strip() for o in options[1:-1].split(",") ]
    table = mydoc.add_table(rows=row, cols=col)
    table.style = 'Table Grid'
    
    #first row include ID and qs
    cell_id = table.cell(0,0)
    cell_id.text = "QN= " + new_mcq.qid

    cell_ques = table.cell(0,1)
    paragraph_cell_ques = cell_ques.paragraphs[0]
    paragraph_cell_ques.text = question.strip()
    is_image = new_mcq.contain_img
    if is_image:
        path ="media/images/"+new_mcq.q_image
        run = paragraph_cell_ques.add_run()
        run.add_picture(path,width=Inches(5), height=Inches(3.2))
        cell_ques.add_paragraph("[file: "+new_mcq.q_image+"]")
        question = "Answer the images question" if question =="" else question

    # add options 

    for i in  range(len(list_op)):
        op_row = table.add_row()
        for j in range (len(op_row.cells)):
            row = op_row.cells[j]
            if j == 0:
                row.text = chr(a + i)+"."
            if j ==1:
                row.text = list_op[i]
    op_row = table.add_row()
    for j in range (len(op_row.cells)):
        row = op_row.cells[j]
        if j == 0:
            row.text = "answer"
        if j ==1:
            row.text = str(new_mcq.answer_q)

    #save to docx 
    print("------- Export ------")
    print("save",str(new_mcq.qid))
    nameDocx = str(subject)+"-"+username+"-"+time+".docx"
    print(nameDocx)
    mydoc.add_paragraph("\n")
    mydoc.save("media/docx/"+nameDocx)

